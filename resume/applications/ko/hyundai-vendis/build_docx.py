"""Build a combined DOCX (resume + self_intro) for Hyundai Vendis application.

Layout:
- Top header: 2-column table — left cell = photo placeholder, right cell = name + contact
- Body: 한 줄 소개, 핵심 역량, 경력, 주요 프로젝트, 학력, 자격
- Page break, then 자기소개서
"""
from pathlib import Path
import re
import sys

import pypandoc
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
RESUME_MD = HERE.parent / "_common" / "resume.md"
SELF_INTRO_MD = HERE / "self_intro.md"
PHOTO = HERE / "photo_resume.jpg"
OUT_MD = HERE / "submission_bundle.md"
OUT_DOCX = HERE / "submission_bundle.docx"

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def load_resume_body() -> tuple[dict, str]:
    """Returns (basic_info_dict, body_markdown_without_basic_info)."""
    text = RESUME_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    cleaned, skip_until_h2 = [], True
    for line in lines:
        if skip_until_h2:
            if line.startswith("## "):
                skip_until_h2 = False
                cleaned.append(line)
            continue
        cleaned.append(line)
    text = "\n".join(cleaned)

    # Extract "## 기본 정보" block
    m = re.search(r"## 기본 정보\n(.*?)(?=\n## )", text, re.DOTALL)
    info_block = m.group(1).strip() if m else ""
    info_lines = [ln.strip().rstrip("  ") for ln in info_block.splitlines() if ln.strip()]
    basic = {
        "name": info_lines[0] if len(info_lines) > 0 else "",
        "title": info_lines[1] if len(info_lines) > 1 else "",
        "rest": info_lines[2:] if len(info_lines) > 2 else [],
    }

    # Remove the 기본 정보 section from body
    body = re.sub(r"## 기본 정보\n.*?(?=\n## )", "", text, flags=re.DOTALL).strip()
    return basic, body


def load_self_intro() -> str:
    text = SELF_INTRO_MD.read_text(encoding="utf-8")
    text = re.sub(r"^# .*\n", "", text, count=1)
    text = re.split(r"\n##\s*5\.", text)[0]
    text = re.sub(r"\n-{3,}\s*$", "", text.rstrip()) + "\n"
    return text


def add_hyperlink(paragraph, url: str, text: str, size_pt: float = 10):
    """Append a hyperlink run to the paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0969DA")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(sz)
    new_run.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text: str, size_pt: float = 10):
    """Add text to paragraph, converting [label](url) to hyperlinks."""
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(size_pt)
        add_hyperlink(paragraph, m.group(2), m.group(1), size_pt=size_pt)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(size_pt)


def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right -> dict(sz, val, color)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side, attrs in kwargs.items():
        el = tc_borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_borders.append(el)
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))


def insert_header_table(doc: Document, basic: dict) -> None:
    """Insert a 2-column header table at the top of doc: [photo | info]."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Set widths
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.5)
    photo_cell, info_cell = table.rows[0].cells
    photo_cell.width = Cm(3.5)
    info_cell.width = Cm(13.5)

    # Photo cell — insert actual image if available, else placeholder
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = photo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if PHOTO.exists():
        run = p.add_run()
        run.add_picture(str(PHOTO), width=Cm(3.5))
    else:
        run = p.add_run("사진\n(3.5cm × 4.5cm)")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x8C, 0x95, 0x9F)
        photo_cell.add_paragraph("")
        border = {"sz": 6, "val": "dashed", "color": "8C959F"}
        set_cell_border(
            photo_cell, top=border, bottom=border, left=border, right=border
        )

    # Info cell — name larger, then title, then contacts
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_name = info_cell.paragraphs[0]
    r_name = p_name.add_run(basic["name"])
    r_name.bold = True
    r_name.font.size = Pt(20)

    if basic.get("title"):
        p_title = info_cell.add_paragraph()
        r_title = p_title.add_run(basic["title"])
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(0x57, 0x60, 0x69)

    if basic.get("rest"):
        p_rest = info_cell.add_paragraph()
        add_text_with_links(p_rest, " · ".join(basic["rest"]), size_pt=10)

    # Move table to the very top (before the first paragraph that pandoc generated)
    body = doc.element.body
    tbl_el = table._tbl
    body.remove(tbl_el)
    body.insert(0, tbl_el)


def main() -> int:
    basic, resume_body = load_resume_body()
    intro_body = load_self_intro()

    combined_md = (
        resume_body
        + "\n\n\\newpage\n\n"
        + "# 자기소개서 - 현대벤디스\n\n"
        + intro_body.lstrip()
    )
    OUT_MD.write_text(combined_md, encoding="utf-8")
    print(f"[built] {OUT_MD}")

    pypandoc.convert_file(
        str(OUT_MD),
        "docx",
        outputfile=str(OUT_DOCX),
        extra_args=["--from=markdown+raw_tex"],
    )

    # Open generated docx and prepend header table
    doc = Document(str(OUT_DOCX))
    insert_header_table(doc, basic)
    doc.save(str(OUT_DOCX))

    print(f"[done] {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
